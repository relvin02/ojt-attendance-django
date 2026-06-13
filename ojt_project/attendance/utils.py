import logging
from django.conf import settings
from io import BytesIO

logger = logging.getLogger(__name__)

def send_sms(to_number: str, body: str) -> bool:
    """Send an SMS to `to_number` with message `body`.

    Priority order:
    1. Semaphore API if `SEMAPHORE_API_KEY` is set in settings.
    2. Twilio if TWILIO_* settings are configured.
    3. Otherwise log a simulated message and return False.
    """
    if not to_number:
        logger.debug("No destination number provided for SMS; skipping.")
        return False

    # semaphore support
    sem_key = getattr(settings, 'SEMAPHORE_API_KEY', None)
    if sem_key:
        try:
            import requests
            payload = {
                "apikey": sem_key,
                "number": to_number,
                "message": body,
            }
            resp = requests.post("https://api.semaphore.co/api/v4/messages", json=payload, timeout=10)
            if resp.ok:
                logger.info("Sent SMS via Semaphore to %s", to_number)
                return True
            else:
                logger.error("Semaphore SMS failed (%s): %s", resp.status_code, resp.text)
        except Exception:
            logger.exception("Exception while sending SMS via Semaphore to %s", to_number)
        # if semaphore fails, fall back to TextBee/Twilio or log below

    # TextBee support (custom provider)
    tb_key = getattr(settings, 'TEXTBEE_API_KEY', None)
    tb_base = getattr(settings, 'TEXTBEE_BASE_URL', 'https://api.textbee.dev')
    tb_device = getattr(settings, 'TEXTBEE_DEVICE_ID', None)
    if tb_key and tb_device:
        try:
            import requests
            headers = {
                "x-api-key": tb_key,
                "Content-Type": "application/json",
            }
            # TextBee expects an array of recipients rather than a single "to" field
            payload = {
                "recipients": [to_number],
                "message": body,
            }
            # use the documented endpoint format
            endpoint = f"{tb_base}/api/v1/gateway/devices/{tb_device}/send-sms"
            resp = requests.post(endpoint, headers=headers, json=payload, timeout=10)
            # some TextBee responses put success flag under data
            body_json = {}
            try:
                body_json = resp.json()
            except Exception:
                pass
            success = body_json.get("success") or body_json.get("data", {}).get("success")
            if resp.ok and success:
                logger.info("Sent SMS via TextBee to %s", to_number)
                return True
            else:
                logger.error("TextBee SMS failed (%s): %s", resp.status_code, resp.text)
        except Exception:
            logger.exception("Exception while sending SMS via TextBee to %s", to_number)
        # if TextBee fails, fall back to Twilio

    # Twilio support
    sid = getattr(settings, 'TWILIO_ACCOUNT_SID', None)
    token = getattr(settings, 'TWILIO_AUTH_TOKEN', None)
    from_number = getattr(settings, 'TWILIO_FROM_NUMBER', None)

    if sid and token and from_number:
        try:
            from twilio.rest import Client
        except Exception as e:
            logger.exception("Twilio library not available: %s", e)
            logger.info("SMS (simulated) to %s: %s", to_number, body)
            return False

        try:
            client = Client(sid, token)
            message = client.messages.create(body=body, from_=from_number, to=to_number)
            logger.info("Sent SMS to %s; sid=%s", to_number, getattr(message, 'sid', ''))
            return True
        except Exception:
            logger.exception("Failed to send SMS via Twilio to %s", to_number)
            return False

    # Fallback: not configured or all providers failed
    logger.info("No SMS provider configured; SMS (logged) to %s: %s", to_number, body)
    return False


def generate_dtr_docx(record, month_name):
    """Generate a DOCX file from DTR record data.
    
    Args:
        record: Dict with 'student', 'dtr_data', 'total_hours', 'total_present', etc.
        month_name: String like "May 2026"
    
    Returns:
        BytesIO object containing the DOCX file
    """
    try:
        from docx import Document
        from docx.shared import Pt, RGBColor, Inches
        from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
    except ImportError:
        logger.error("python-docx not installed")
        return None
    
    try:
        doc = Document()
        
        # Title
        title = doc.add_paragraph()
        title_run = title.add_run('Daily Time Record (DTR)')
        title_run.bold = True
        title_run.font.size = Pt(16)
        title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        
        # Month
        month_para = doc.add_paragraph(month_name)
        month_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        
        # Student Info
        doc.add_paragraph()  # spacing
        student = record['student']
        doc.add_paragraph(f"ID Number: {student.id_number}")
        doc.add_paragraph(f"Name: {student.first_name} {student.last_name}")
        doc.add_paragraph(f"Batch: {student.batch.name if student.batch else 'N/A'}")
        if student.company:
            doc.add_paragraph(f"Company: {student.company.name}")
        
        # DTR Table
        doc.add_paragraph()  # spacing
        table = doc.add_table(rows=1, cols=6)
        table.style = 'Light Grid Accent 1'
        
        # Header row
        header_cells = table.rows[0].cells
        headers = ['Date', 'Day', 'Check-In', 'Check-Out', 'Hours', 'Status']
        for i, header in enumerate(headers):
            header_cells[i].text = header
        
        # Data rows
        for day in record['dtr_data']:
            row_cells = table.add_row().cells
            row_cells[0].text = day['date'].strftime('%m/%d/%Y')
            row_cells[1].text = day['day'][:3]  # Mon, Tue, etc.
            row_cells[2].text = day['check_in'].strftime('%H:%M') if day['check_in'] else '--'
            row_cells[3].text = day['check_out'].strftime('%H:%M') if day['check_out'] else '--'
            row_cells[4].text = f"{day['hours']}h" if day['hours'] > 0 else '--'
            
            status = day['status']
            if status == 'present':
                status_text = 'Present'
                if day['was_late']:
                    status_text += ' (Late)'
            elif status == 'no_record':
                status_text = 'No Record'
            else:
                status_text = status.capitalize()
            row_cells[5].text = status_text
        
        # Summary
        doc.add_paragraph()  # spacing
        summary_para = doc.add_paragraph()
        summary_para.add_run('Summary:\n').bold = True
        summary_para.add_run(f"Total Hours: {record['total_hours']}\n")
        summary_para.add_run(f"Days Present: {record['total_present']}\n")
        summary_para.add_run(f"Days Late: {record['total_late']}\n")
        summary_para.add_run(f"Days Absent: {record['total_absent']}\n")
        
        # Signature lines
        doc.add_paragraph()
        doc.add_paragraph()
        sig_table = doc.add_table(rows=2, cols=3)
        sig_table.autofit = False
        
        sig_labels = ['Prepared by', 'Checked by', 'Approved by']
        for i, label in enumerate(sig_labels):
            sig_table.rows[1].cells[i].text = label
        
        # Save to BytesIO
        output = BytesIO()
        doc.save(output)
        output.seek(0)
        return output.getvalue()
    
    except Exception as e:
        logger.exception(f"Error generating DOCX: {e}")
        return None

